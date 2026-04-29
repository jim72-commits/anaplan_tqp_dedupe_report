"""Generates the SRB Overlap Report user guide as a .docx.

Run from the repo root (or this directory) with:

    python docs/build_user_guide.py

Produces:  docs/Overlap_Report_User_Guide.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

OUT_PATH = Path(__file__).resolve().parent / "Overlap_Report_User_Guide.docx"

# Moody's-aligned palette used in the report itself.
MOODY_RED = RGBColor(0xE3, 0x06, 0x13)
MOODY_BLUE = RGBColor(0x12, 0x35, 0x6B)
TEXT_DARK = RGBColor(0x1A, 0x1F, 0x2C)
TEXT_MUTED = RGBColor(0x5C, 0x65, 0x75)
PALE_GREY = RGBColor(0xF4, 0xF5, 0xF7)
RULE_GREY = RGBColor(0xD0, 0xD4, 0xDB)


def shade_cell(cell, hex_fill: str) -> None:
    """Apply a background fill colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "D0D4DB", size: str = "4") -> None:
    """Light grey borders around a single cell (used for tables)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_horizontal_rule(doc: Document, color: RGBColor = RULE_GREY) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_heading(paragraph, color: RGBColor, size_pt: int) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"


def add_callout(doc: Document, label: str, body: str, label_color: RGBColor = MOODY_BLUE) -> None:
    """A two-row mini-table styled as a tip / note callout."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.3)
    shade_cell(cell, "F4F5F7")
    set_cell_borders(cell, color="D0D4DB")
    p = cell.paragraphs[0]
    label_run = p.add_run(label + "  ")
    label_run.bold = True
    label_run.font.color.rgb = label_color
    label_run.font.size = Pt(10)
    body_run = p.add_run(body)
    body_run.font.color.rgb = TEXT_DARK
    body_run.font.size = Pt(10)
    doc.add_paragraph()


def add_kv_table(doc: Document, header: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    """Two-column reference table (key + description)."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.3)

    header_cells = table.rows[0].cells
    for idx, label in enumerate(header):
        header_cells[idx].width = Inches(2.0 if idx == 0 else 4.3)
        shade_cell(header_cells[idx], "1A1F2C")
        set_cell_borders(header_cells[idx], color="1A1F2C")
        p = header_cells[idx].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, (key, desc) in enumerate(rows, start=1):
        row_cells = table.rows[r_idx].cells
        row_cells[0].width = Inches(2.0)
        row_cells[1].width = Inches(4.3)
        if r_idx % 2 == 0:
            shade_cell(row_cells[0], "F4F5F7")
            shade_cell(row_cells[1], "F4F5F7")
        set_cell_borders(row_cells[0])
        set_cell_borders(row_cells[1])

        kp = row_cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(10)
        kr.font.color.rgb = TEXT_DARK

        dp = row_cells[1].paragraphs[0]
        dr = dp.add_run(desc)
        dr.font.size = Pt(10)
        dr.font.color.rgb = TEXT_DARK

    doc.add_paragraph()


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_DARK


def body_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_DARK


def heading(doc: Document, text: str, level: int) -> None:
    """Apply consistent typography on top of the built-in heading styles."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    color, size = {
        1: (MOODY_RED, 22),
        2: (MOODY_BLUE, 16),
        3: (TEXT_DARK, 13),
    }[level]
    style_heading(h, color, size)


def title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Overlap Report")
    title_run.font.size = Pt(34)
    title_run.bold = True
    title_run.font.color.rgb = MOODY_RED
    title_run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("SRB Assignment — User Guide")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = MOODY_BLUE
    sub_run.font.name = "Calibri"

    blurb = doc.add_paragraph()
    bl = blurb.add_run(
        "How to load source data, read the views, drill into shared MixedIDs, "
        "and export findings."
    )
    bl.font.size = Pt(11)
    bl.italic = True
    bl.font.color.rgb = TEXT_MUTED

    add_horizontal_rule(doc)


def build() -> Path:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_block(doc)

    # ------------------------------------------------------------------
    heading(doc, "1. What the report does", level=1)
    body_para(
        doc,
        "The Overlap Report compares two source files — your SRB assignment "
        "data and the Territory Rules export — and surfaces every territory / "
        "rule pair that shares MixedIDs with another rule. For each overlap, "
        "the report tells you which rules collide, how many MixedIDs are "
        "duplicated, what parent / child accounts are involved, and how much "
        "SRB revenue sits inside the overlap.",
    )
    body_para(
        doc,
        "Use it before locking territory plans, when reconciling "
        "rule changes, or whenever a Sales Planner asks why two reps are "
        "claiming the same account.",
    )

    add_callout(
        doc,
        "Privacy",
        "All processing happens locally in your browser. Source files never "
        "leave your machine — there's no upload to a server.",
    )

    # ------------------------------------------------------------------
    heading(doc, "2. Getting started", level=1)

    heading(doc, "2.1 Open the report", level=2)
    add_steps(
        doc,
        [
            "Open the index.html in a modern browser (Chrome, Edge, or Firefox).",
            "Wait a moment for the application shell to load. The first time you "
            "open the page you'll land on the Coin Sort tab — click "
            "“Overlap Report” in the top tab strip to switch.",
        ],
    )

    heading(doc, "2.2 What you'll see on first load", level=2)
    add_bullets(
        doc,
        [
            "Header strip with the report title, two empty drop zones (SRB "
            "Source Data and Territory Rules), the Analyze button, an Imported "
            "timestamp pill, and a New Files button.",
            "Toolbar with search, four multi-select filter chips, view switcher "
            "(Grid / Matrix / Heat Map), Export Overlap, Export Details, and "
            "the Views (saved-view) menu.",
            "Empty grid prompting you to drop source files.",
        ],
    )

    # ------------------------------------------------------------------
    heading(doc, "3. Loading data", level=1)

    heading(doc, "3.1 Required source files", level=2)
    body_para(
        doc,
        "The report consumes two CSV files. Both should be UTF-8 with a "
        "header row. Column order does not matter — the parser matches by name.",
    )

    add_kv_table(
        doc,
        ("SRB Source Data — column", "Notes"),
        [
            ("MIXED ID", "Required. Unique ID of the assignment row."),
            ("Territory Code", "Required. Maps the row to a territory."),
            ("Territory Name", "Recommended. Used for friendly labels."),
            ("Rule Code", "Required. Maps the row to a rule."),
            ("Rule Name", "Recommended. Used for friendly labels."),
            ("Associated Job Type", "Filterable in the report."),
            ("Sales Team / Sales Sub Team", "Both filterable."),
            ("Parent Account Code / Parent Account Name", "Shown in detail drawer and exports."),
            ("Account ID / Account Name", "Shown in detail drawer and exports."),
            ("Country", "Drives the Country Heat Map view."),
            ("20XX SRB $ (numeric)", "Optional. If present, the report shows SRB-weighted overlaps and dollar totals."),
        ],
    )

    add_kv_table(
        doc,
        ("Territory Rules (SYS09B) — column", "Notes"),
        [
            ("Rule Code", "Required. Joins back to the SRB data."),
            ("Territory Code", "Required."),
            ("TERRITORY NAME", "Used as the friendly label everywhere."),
            ("TERRITORY STATUS", "Visible in the rule drawer."),
            ("Region / Sub-Region / Country / State / City", "Used for filters and for hover detail."),
            ("Segment / Sub-Segment / Parent Segment", "Used in the rule signature for overlap detection."),
            ("Tiers, Strategic Account, Named Accounts Only", "Used in overlap signature."),
            ("Product Family / Product Group", "Used in overlap signature."),
            ("SRB / WSR (numeric)", "Shown alongside rule details."),
            ("RULE NOTES", "Free text shown in the rule drawer."),
        ],
    )

    heading(doc, "3.2 Drop, click, and analyze", level=2)
    add_steps(
        doc,
        [
            "Drag the SRB Source Data CSV onto the first drop zone, or click the "
            "zone to browse for it.",
            "Drag the Territory Rules CSV onto the second drop zone (or click "
            "to browse).",
            "When both zones are filled, the Analyze button activates. Click it "
            "to compute overlaps. Large files may take a few seconds — you'll "
            "see a progress message in the header.",
            "When analysis finishes, the Imported pill timestamps the run and "
            "the grid populates.",
        ],
    )

    add_callout(
        doc,
        "Tip",
        "After a successful analysis, the parsed dataset is cached locally. "
        "The next time you open the page, both drop zones say “Cached” and "
        "the previous run is restored automatically. Click New Files to clear "
        "the cache and start over.",
    )

    # ------------------------------------------------------------------
    heading(doc, "4. The Overlap Report tab", level=1)

    heading(doc, "4.1 Toolbar", level=2)
    add_kv_table(
        doc,
        ("Control", "Purpose"),
        [
            ("Search box", "Free-text match on rule code, rule name, territory code, or territory name. Combined with all filters (AND)."),
            ("All Sales Teams ▾", "Multi-select chip. Tick one or more values to narrow the grid; click “Clear” inside the dropdown to remove all."),
            ("All Sales Sub Teams ▾", "As above for sub-team."),
            ("All Job Types ▾", "As above for the Associated Job Type column."),
            ("All Territories ▾", "As above for territory codes/names."),
            ("✕ Clear", "Wipes search and every filter chip in one click."),
            ("View switcher", "Grid / Matrix / Heat Map. Same data, different lens — see 4.2 onwards."),
            ("Export Overlap", "Top-level CSV: one row per rule with totals. Always reflects the current filter set."),
            ("Export Details", "MixedID-level CSV: one row per shared MixedID per overlapping rule pair."),
            ("★ Views", "Save and reload combinations of search + filters + sort — see section 7."),
        ],
    )

    heading(doc, "4.2 Grid view (default)", level=2)
    body_para(
        doc,
        "The grid lists every rule that has at least one overlap with another "
        "rule. Each row is sortable — click any column header to toggle "
        "ascending / descending. Heat shading on the row reflects severity "
        "(deeper red = more shared MixedIDs).",
    )
    add_kv_table(
        doc,
        ("Column", "What it shows"),
        [
            ("☆ (pin)", "Click to add the row to the Watchlist banner. Pinned rows stay visible when you filter or sort and persist across reloads."),
            ("Territory Name", "Friendly territory label (full territory + rule code shown on hover)."),
            ("Rule Name", "Friendly rule label (full code on hover)."),
            ("Job Type", "Distinct job types attached to the rule."),
            ("Overlapped Rules", "Count of OTHER rules sharing at least one MixedID with this row."),
            ("Overlapped Parent Acct", "Distinct parent accounts that appear in any overlap."),
            ("Overlapped Child Acct", "Distinct child accounts that appear in any overlap."),
            ("Overlapped MixedID", "Distinct MixedIDs shared with one or more other rules. Sparkbar shows relative severity."),
            ("Overlapped SRB $", "SRB dollars contained inside the overlap (only when the source CSV had an SRB column)."),
        ],
    )

    heading(doc, "4.3 Pin / Watchlist", level=2)
    body_para(
        doc,
        "Click the star icon at the start of any row to pin it. Pinned rows "
        "appear in the Watchlist bar above the grid and remain pinned across "
        "filtering, sorting, and page reloads. Click the × on a chip in the "
        "Watchlist bar to remove the pin.",
    )

    heading(doc, "4.4 Compare mode", level=2)
    add_steps(
        doc,
        [
            "Tick the “SELECT TO COMPARE” checkbox in the grid's panel header. "
            "A column of compare checkboxes appears beside each row.",
            "Tick exactly two rows. The Compare modal opens automatically with "
            "the two rules side by side, showing matching rule attributes, "
            "shared MixedIDs, and synchronized scrolling between the two "
            "panels.",
            "Close the modal to return to the grid; the compare checkboxes "
            "remain ticked so you can swap pairs quickly.",
        ],
    )

    heading(doc, "4.5 Quick-rank chips", level=2)
    body_para(
        doc,
        "“Top 10 by SRB $” and “Top 10 by MixedIDs” are shortcut filters that "
        "shrink the grid to the 10 rules with the most overlap exposure on "
        "that metric. Click the chip again to clear.",
    )

    heading(doc, "4.6 Matrix view", level=2)
    body_para(
        doc,
        "A bilateral rule × rule heatmap. Both axes list the same rules "
        "ordered by severity score. Each cell encodes the number of MixedIDs "
        "the row-rule and column-rule share — colour intensity follows the "
        "low → high gradient at the top of the panel.",
    )
    add_bullets(
        doc,
        [
            "Hover over a cell to see the exact shared-MixedID count and the two rule labels.",
            "Click a cell to open the side-by-side compare modal for that rule pair.",
            "Row and column labels truncate with “…” when they don't fit; the full label is always available on hover.",
            "The matrix re-renders on filter changes and on window resize, so it always reflects the current scope.",
        ],
    )

    heading(doc, "4.7 Country Heat Map view", level=2)
    body_para(
        doc,
        "An embedded world map with one bubble per country, sized and "
        "coloured by the active metric. Useful for spotting geographic "
        "concentrations of overlap.",
    )
    add_kv_table(
        doc,
        ("Control", "Purpose"),
        [
            ("Color by: MixedIDs / SRB $ / Rules", "Switches the metric driving bubble size & colour."),
            ("Zoom in / Zoom out / Reset", "Geographic zoom for inspecting dense regions."),
            ("Click a bubble", "Filters every other view to that country. A “🌐 Country: X” bar appears at the top of the page; click its × to clear."),
        ],
    )

    add_callout(
        doc,
        "Note",
        "Country values that don't resolve to an ISO country (typos, "
        "test data, internal codes) are logged to the browser console and "
        "left out of the map. Run with the developer console open if you "
        "want to see exactly which values were skipped.",
    )

    # ------------------------------------------------------------------
    heading(doc, "5. Drilling into details", level=1)

    heading(doc, "5.1 Open the detail drawer", level=2)
    body_para(
        doc,
        "Click anywhere on a grid row (outside the pin / compare controls) "
        "to open the Overlap Detail drawer at the bottom of the screen. "
        "Click the same row again, press Esc, or click the × on the drawer "
        "to close it.",
    )

    heading(doc, "5.2 Reading the drawer", level=2)
    add_bullets(
        doc,
        [
            "Header summary strip: Territory Name · Rule Name · MixedIDs · Parent Accts · Child Accts. Hover the names to reveal the original codes.",
            "Detail grid: one row per overlapping rule, with overlapped territory, rule, job type, parent / child / MixedID counts, and overlapped SRB $.",
            "Sparkbar in the Overlapped MixedID column shows that row's overlap relative to the largest overlap in the drawer.",
            "Click the ▼ icon on any drawer row to expand it inline and see the actual shared MixedIDs (with parent / child / SRB / job type / sales team).",
        ],
    )

    heading(doc, "5.3 Export from inside the drawer", level=2)
    add_bullets(
        doc,
        [
            "Export Shared MixedIDs (top of drawer) — every shared MixedID for the focused rule across all of its overlaps, deduplicated.",
            "▼ row export — the shared MixedIDs for just that single overlapping pair.",
        ],
    )

    # ------------------------------------------------------------------
    heading(doc, "6. Filtering", level=1)
    body_para(
        doc,
        "Filters combine with AND across types, OR within a single multi-select. "
        "Everything else in the page (KPIs in the watchlist, Matrix, Heat Map, "
        "exports) updates with the filter set instantly.",
    )
    add_bullets(
        doc,
        [
            "Search: matches rule code, rule name, territory code, or territory name (case-insensitive, contains).",
            "Sales Teams / Sub Teams / Job Types / Territories: tick one or more values inside each chip's dropdown.",
            "Use “Clear” inside an open dropdown to wipe just that filter, or the toolbar ✕ Clear button to reset everything.",
            "Click a country bubble in the Heat Map to add a country filter; clear it from the country bar above the grid.",
        ],
    )

    # ------------------------------------------------------------------
    heading(doc, "7. Saved views", level=1)
    body_para(
        doc,
        "The ★ Views menu stores combinations of search text + filter chips + "
        "sort, named by you, in the browser's local storage.",
    )
    add_steps(
        doc,
        [
            "Set up the search / filters / sort you want to save.",
            "Open the ★ Views menu and click “Save current view…”.",
            "Type a name (e.g. “EMEA Banking — overlaps over 5 MixedIDs”) and confirm.",
            "Reload at any time by re-opening the menu and clicking the saved name.",
            "Hover a saved view to reveal a delete button.",
        ],
    )

    # ------------------------------------------------------------------
    heading(doc, "8. Exporting", level=1)

    add_kv_table(
        doc,
        ("Export", "Output"),
        [
            (
                "Export Overlap",
                "Overlap_Report.csv — one row per rule with overlap counts, "
                "Overlapped SRB $, the list of overlapped rule codes, and the "
                "matching list of friendly “Territory - Rule” labels.",
            ),
            (
                "Export Details",
                "Overlap_Detail_MixedID.csv — one row per (source rule × "
                "overlapping rule × shared MixedID), with parent / child "
                "account, MixedID, SRB $, job type, and sales team.",
            ),
            (
                "Export Shared MixedIDs (drawer header)",
                "All_Shared_MixedIDs_Rule_<rule>.csv — every shared MixedID "
                "for the focused rule, deduplicated, with overlapping rules "
                "list per row.",
            ),
            (
                "Export Shared MixedIDs (drawer row ▼)",
                "Shared_MixedIDs_<ruleA>_vs_<ruleB>.csv — the shared MixedIDs "
                "for that single rule pair.",
            ),
        ],
    )
    add_callout(
        doc,
        "Tip",
        "Every export respects the current search + filter + view scope. "
        "If you've selected a country on the Heat Map, exports include only "
        "rows that fall in that country.",
    )

    # ------------------------------------------------------------------
    heading(doc, "9. Coin Sort tab", level=1)
    body_para(
        doc,
        "The Coin Sort tab is a separate engine for replaying rule logic "
        "against the SRB data and producing the canonical territory "
        "assignment for each MixedID. It uses the same two source files as "
        "the Overlap Report.",
    )

    heading(doc, "9.1 Single Territory mode", level=2)
    add_steps(
        doc,
        [
            "Tick one or more territories on the left grid (use the search "
            "box and All / None buttons to scope quickly).",
            "Tick the rules you want to evaluate on the right grid.",
            "Click Run Coin Sort. Results appear in the lower panel and can be exported.",
        ],
    )

    heading(doc, "9.2 Bulk Refresh mode", level=2)
    body_para(
        doc,
        "Runs Coin Sort across the full rule set. Click Run Coin Sort, watch "
        "the progress strip, and use Export Results CSV when the run "
        "completes.",
    )

    # ------------------------------------------------------------------
    heading(doc, "10. Keyboard shortcuts", level=1)
    add_kv_table(
        doc,
        ("Shortcut", "Action"),
        [
            ("Ctrl + K", "Open the command palette (jump to any view, filter, or saved view)."),
            ("/", "Focus the search box."),
            ("G", "Switch to Grid view."),
            ("M", "Switch to Matrix view."),
            ("C", "Switch to Country Heat Map view."),
            ("Shift + E", "Toggle table density (compact ↔ comfortable)."),
            ("↑ / ↓", "Move row selection up / down in the grid."),
            ("Esc", "Close any open dropdown, modal, or drawer."),
            ("?", "Open the Keyboard Shortcuts overlay."),
            ("Ctrl + P", "Print or save as PDF."),
        ],
    )

    # ------------------------------------------------------------------
    heading(doc, "11. Tips & troubleshooting", level=1)
    add_bullets(
        doc,
        [
            "Cached vs fresh data: if drop-zones show “Cached”, the page is "
            "using the previous run. Click New Files (top right of header) "
            "to clear the cache and re-import.",
            "Dropdowns won't close: click anywhere outside the chip or press "
            "Esc.",
            "A row I expect is missing: confirm the rule code appears in BOTH "
            "files, and that none of the active filters exclude it. The "
            "Watchlist bar is independent of filters and is a quick way to "
            "verify a row exists.",
            "The Heat Map says “1 country value(s) not on map”: typically "
            "caused by typos or non-ISO codes in the Country column. The "
            "browser console lists the offending values.",
            "Performance on very large files: the parser streams the SRB "
            "file, but extreme datasets (>2M rows) may take several seconds. "
            "Filtering is fully client-side and fast once parsing finishes.",
            "Sharing results: use Export Overlap or Export Details and "
            "attach the resulting CSV — the rendered page itself is not "
            "designed to be re-hosted from elsewhere.",
        ],
    )

    add_callout(
        doc,
        "Need help?",
        "Press Ctrl + K to open the command palette, or ? to open the "
        "Keyboard Shortcuts overlay. Both expose every action in the report "
        "and are searchable.",
        label_color=MOODY_RED,
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build()
    print(f"Wrote: {out}")
